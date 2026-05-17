#!/usr/bin/env python3
from __future__ import annotations

import json
import platform
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path("/Users/sdf/Desktop/论文复现")
PACK = ROOT / "reproduction_pack"
ACL_DIR = PACK / "acl2023_interpretable_definition_generation"
EACL_DIR = PACK / "eacl2024_semantic_change_benchmark"


def set_east_asia_font(run, font_name: str = "Songti SC") -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(document: Document, text: str = "", bold: bool = False) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    set_east_asia_font(run)


def add_codeblock(document: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "Courier New")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def git_head(path: Path) -> str:
    return subprocess.check_output(
        ["git", "-C", str(path), "rev-parse", "HEAD"], text=True
    ).strip()


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def main() -> None:
    acl_generated = pd.read_csv(
        ACL_DIR / "artifacts" / "semantic_change_case" / "generated_definitions.tsv",
        sep="\t",
    )
    acl_labels = pd.read_csv(
        ACL_DIR / "artifacts" / "semantic_change_case" / "sense_labels.tsv",
        sep="\t",
    )
    codwoe_generated = pd.read_csv(
        ACL_DIR / "artifacts" / "codwoe_en_trial" / "en_trial_generated.tsv",
        sep="\t",
    )
    codwoe_metrics = pd.read_csv(
        ACL_DIR / "artifacts" / "codwoe_en_trial" / "en_trial_metrics.tsv",
        sep="\t",
        header=None,
        names=["metric", "score"],
    )

    eacl_preds = pd.read_csv(EACL_DIR / "artifacts" / "predictions.csv", sep="\t")
    eacl_result = json.loads(read_text(EACL_DIR / "artifacts" / "result.json"))
    eacl_labels = pd.read_csv(
        EACL_DIR / "artifacts" / "benchmark_labels_subset.tsv", sep="\t"
    )

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("两篇语义变化论文本地复现说明文档")
    set_east_asia_font(run)
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("工作目录：/Users/sdf/Desktop/论文复现")
    set_east_asia_font(run)

    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "文档目的：按照老师要求，将两篇论文的本地复现情况、实际运行的数据、输入输出、基准数据、命令、结果与限制条件集中整理为可提交的 docx 说明文档。",
    )
    add_paragraph(
        doc,
        "总体结论：两篇论文的官方代码都已经在本机跑通。第 1 篇（ACL 2023）不仅完成了“真实语义变化数据 -> 定义生成 -> 原型 sense label”的核心流程复现，还补充跑通了 CoDWoE English trial 的金标评测层；第 2 篇（EACL 2024）完成了官方 benchmark 框架在真实数据子集上的可运行复现，并得到可记录的 Spearman 结果。",
    )

    doc.add_heading("1. 复现环境", level=1)
    add_paragraph(doc, f"操作系统：{platform.platform()}")
    add_paragraph(doc, f"机器架构：{platform.machine()}")
    add_paragraph(doc, "ACL 2023 环境：/Users/sdf/Desktop/论文复现/.conda/defgen")
    add_paragraph(doc, "EACL 2024 环境：/Users/sdf/Desktop/论文复现/.conda/lscd")
    add_paragraph(doc, f"ACL 2023 本地代码提交：{git_head(ROOT / 'definition_modeling')}")
    add_paragraph(doc, f"EACL 2024 本地代码提交：{git_head(ROOT / 'LSCDBenchmark')}")

    doc.add_heading("2. 论文一：ACL 2023", level=1)
    add_paragraph(
        doc,
        "论文标题：Interpretable Word Sense Representations via Definition Generation: The Case of Semantic Change Analysis",
    )
    add_paragraph(doc, "本地代码目录：/Users/sdf/Desktop/论文复现/definition_modeling")
    add_paragraph(
        doc,
        "本次复现目标：不只验证脚本能否运行，而是尽量贴近论文的核心方法，即在语义变化数据上为每条 usage 自动生成定义，再从同一 sense cluster 中选出最原型的定义作为 sense label；同时再向论文原始 definition-generation benchmark 推进一步，补充一层带金标的 CoDWoE 英语评测。",
    )

    doc.add_heading("2.1 实际使用的数据", level=2)
    add_paragraph(
        doc,
        "原论文仓库 README 提到的定义建模数据包括 WordNet、Oxford 和 CoDWoE。本次本地复现采用两层 ACL 数据：第一层直接贴近论文 semantic change analysis 主线，第二层向原始 definition-generation benchmark 推进一步。",
    )
    add_paragraph(doc, "第一层：真实语义变化案例。")
    add_paragraph(doc, "源数据根目录：/Users/sdf/Desktop/论文复现/LSCDBenchmark/wug/nor_dia_change/subset1")
    add_paragraph(doc, "选取词项：kjemi（基本不变）、damp（中等变化）、plattform（强变化）")
    add_paragraph(doc, "准备后的输入文件：/Users/sdf/Desktop/论文复现/reproduction_pack/acl2023_interpretable_definition_generation/artifacts/semantic_change_case/semantic_change_subset.tsv")
    add_paragraph(doc, f"输入 usage 数：{len(acl_generated)}")
    add_paragraph(
        doc,
        "输入文件包含字段：id, word, pos, date, period, cluster, target_indices, Targets, Context。",
    )
    add_paragraph(doc, "第二层：CoDWoE English trial 金标评测。")
    add_paragraph(doc, "源数据文件：/Users/sdf/Desktop/论文复现/acl2023_benchmarks/codwoe/trial_data/en.trial.complete.json")
    add_paragraph(doc, "准备后的输入文件：/Users/sdf/Desktop/论文复现/reproduction_pack/acl2023_interpretable_definition_generation/artifacts/codwoe_en_trial/en_trial_input.tsv")
    add_paragraph(doc, f"输入样本数：{len(codwoe_generated)}")
    add_paragraph(
        doc,
        "输入文件包含字段：id, Targets, Context, Definition, POS, Type。其中 Definition 是 trial split 自带的金标 gloss，用于自动评测。",
    )

    doc.add_heading("2.2 实际运行命令", level=2)
    add_paragraph(doc, "真实语义变化主流程脚本：")
    add_codeblock(
        doc,
        read_text(
            ACL_DIR / "commands" / "run_semantic_change_pipeline.sh"
        ),
    )
    add_paragraph(doc, "CoDWoE English trial benchmark 脚本：")
    add_codeblock(
        doc,
        read_text(
            ACL_DIR / "commands" / "run_codwoe_trial_benchmark.sh"
        ),
    )

    doc.add_heading("2.3 输入与输出", level=2)
    add_paragraph(doc, "语义变化案例输出目录：/Users/sdf/Desktop/论文复现/reproduction_pack/acl2023_interpretable_definition_generation/artifacts/semantic_change_case")
    add_paragraph(doc, "关键输出文件：")
    add_paragraph(doc, "1. semantic_change_subset.tsv：准备后的真实输入数据")
    add_paragraph(doc, "2. generated_definitions.tsv：每条 usage 的自动生成定义")
    add_paragraph(doc, "3. definition_embeddings.npz：定义嵌入")
    add_paragraph(doc, "4. sense_labels.tsv：每个 cluster 的原型定义标签")
    add_paragraph(doc, "CoDWoE benchmark 输出目录：/Users/sdf/Desktop/论文复现/reproduction_pack/acl2023_interpretable_definition_generation/artifacts/codwoe_en_trial")
    add_paragraph(doc, "5. en_trial_input.tsv：由官方 json 转成的本地评测输入")
    add_paragraph(doc, "6. en_trial_generated.tsv：200 条样本的生成定义结果")
    add_paragraph(doc, "7. en_trial_metrics.tsv：自动评测分数")

    doc.add_heading("2.4 真实语义变化链路结果", level=2)
    sample_rows = []
    for _, row in acl_generated[["word", "cluster", "Generated_Definition"]].head(6).iterrows():
        sample_rows.append(
            [
                str(row["word"]),
                str(row["cluster"]),
                str(row["Generated_Definition"])[:80],
            ]
        )
    add_table(doc, ["词", "聚类编号", "生成定义（截断显示）"], sample_rows)

    add_paragraph(doc, "聚类原型标签输出如下：")
    label_rows = []
    for _, row in acl_labels.iterrows():
        label_rows.append(
            [
                str(row["Targets"]),
                str(row["Clusters"]),
                str(row["Definitions"])[:90],
            ]
        )
    add_table(doc, ["词", "聚类", "原型定义标签"], label_rows)

    add_paragraph(
        doc,
        "结果解读：kjemi 只保留了 1 个 cluster，因此最终只有 1 个原型定义。damp 与 plattform 存在多个 cluster，脚本能够输出每个 cluster 的原型定义；但若某个 cluster 中样本数少于 3，脚本会直接输出“Too few examples to generate a proper definition!”，这属于原仓库脚本的设计逻辑，不是本次运行错误。",
    )
    add_paragraph(
        doc,
        "这条链路的复现意义在于：它直接对应论文题目中的 semantic change analysis 应用部分，说明本地环境已经可以跑出“usage -> definition -> prototype sense label”的解释性主流程。",
    )

    doc.add_heading("2.5 CoDWoE English trial 金标评测", level=2)
    add_paragraph(
        doc,
        "为了让 ACL 2023 不只停留在应用案例，本次又补充跑了仓库 README 提到的 CoDWoE 路线中的 English trial。该 split 自带金标 gloss，因此可以对生成结果做自动评测。",
    )
    metric_rows = []
    for _, row in codwoe_metrics.iterrows():
        metric_rows.append([str(row["metric"]), f"{float(row['score']):.4f}"])
    add_table(doc, ["指标", "分数"], metric_rows)
    codwoe_rows = []
    for _, row in codwoe_generated[["Targets", "Definition", "Generated_Definition"]].head(5).iterrows():
        codwoe_rows.append(
            [
                str(row["Targets"]),
                str(row["Definition"])[:70],
                str(row["Generated_Definition"])[:70],
            ]
        )
    add_table(doc, ["词", "金标定义（截断）", "生成定义（截断）"], codwoe_rows)
    add_paragraph(
        doc,
        "结果解读：CoDWoE English trial 一共 200 条样本，本次得到 sacrebleu=6.5549、rougeL=0.2190、exact_match=0.0000。这个分数不能代表完整论文 benchmark 的最终水平，但它证明了该仓库的定义生成模型已经能在带金标的 benchmark split 上完成端到端评测。",
    )
    add_paragraph(
        doc,
        "复现判断：ACL 2023 现在既有真实语义变化应用链路，也有一层带金标的 CoDWoE benchmark 评测，因此比单纯案例演示更接近“基本上复现论文”的要求；但还不能宣称已经完整覆盖 WordNet、Oxford 和完整 CoDWoE 全套大规模实验。",
    )

    doc.add_heading("3. 论文二：EACL 2024", level=1)
    add_paragraph(doc, "论文标题：Computational modeling of semantic change")
    add_paragraph(doc, "本地代码目录：/Users/sdf/Desktop/论文复现/LSCDBenchmark")
    add_paragraph(
        doc,
        "论文性质说明：这篇文章本身是 tutorial / benchmark 论文，不是单一模型论文。因此“复现论文”在实践上对应为：使用作者给出的 benchmark 框架，在官方支持的数据集与任务配置上跑出可复核结果。",
    )

    doc.add_heading("3.1 实际使用的数据与基准配置", level=2)
    add_paragraph(doc, "数据集配置：nordiachange_1")
    add_paragraph(doc, "实际数据根目录：/Users/sdf/Desktop/论文复现/LSCDBenchmark/wug/nor_dia_change/subset1")
    add_paragraph(doc, "选取 5 个词作为本地子集：kjemi、egg、damp、fil、plattform")
    add_paragraph(doc, "任务：lscd_graded")
    add_paragraph(doc, "模型：apd_compare_all")
    add_paragraph(doc, "WiC 模型：contextual_embedder")
    add_paragraph(doc, "相似度：cosine")
    add_paragraph(doc, "Transformer 检查点：prajjwal1/bert-tiny")
    add_paragraph(
        doc,
        "之所以使用较小的 checkpoint，是因为本地复现目标是“验证 benchmark 主链路并得到可对照结果”，不是在当前机器上追求整篇教程中所有模型的最优成绩。",
    )

    doc.add_heading("3.2 实际运行命令", level=2)
    add_codeblock(
        doc,
        read_text(EACL_DIR / "commands" / "run_verified_subset.sh"),
    )

    doc.add_heading("3.3 基准标签与输出", level=2)
    add_paragraph(doc, "基准标签文件：/Users/sdf/Desktop/论文复现/reproduction_pack/eacl2024_semantic_change_benchmark/artifacts/benchmark_labels_subset.tsv")
    label_rows = []
    for _, row in (
        eacl_labels[["lemma", "change_graded"]]
        .sort_values("lemma")
        .iterrows()
    ):
        label_rows.append([str(row["lemma"]), f"{row['change_graded']:.6f}"])
    add_table(doc, ["词", "金标 change_graded"], label_rows)

    add_paragraph(doc, "实际输出文件：")
    add_paragraph(doc, "1. /Users/sdf/Desktop/论文复现/reproduction_pack/eacl2024_semantic_change_benchmark/artifacts/result.json")
    add_paragraph(doc, "2. /Users/sdf/Desktop/论文复现/reproduction_pack/eacl2024_semantic_change_benchmark/artifacts/predictions.csv")
    add_paragraph(doc, f"本次子集复现的 Spearman 结果：{eacl_result['score']}")

    pred_rows = []
    for _, row in eacl_preds.iterrows():
        pred_rows.append(
            [
                str(row["instance"]),
                f"{row['prediction']:.6f}",
                f"{row['label']:.6f}",
            ]
        )
    add_table(doc, ["词", "预测值", "金标"], pred_rows)

    add_paragraph(
        doc,
        "结果解读：5 个词子集上的 Spearman 为 0.3591，说明在当前轻量配置下，benchmark 主流程可以稳定完成，但模型能力与词集规模还不足以逼近论文中更全面、更强模型配置下的表现。也正因此，这个结果更适合作为“本地基本复现”而不是“完全复现全部实验结论”。",
    )

    doc.add_heading("4. 本次复现中的代码兼容性修改", level=1)
    add_paragraph(
        doc,
        "ACL 2023 仓库修改：",
    )
    add_paragraph(doc, "1. generate_t5.py：将 sampling/filter 的 0/1 参数显式转成布尔值，避免当前 transformers 下 generation 返回 None。")
    add_paragraph(doc, "2. embed_definitions.py、sense_label.py、extract_usage_embeddings.py、definition_pair_similarity.py：将旧接口 eval_metrics() 改为 eval()。")
    add_paragraph(doc, "3. evaluate_simple.py：改为只按实际请求加载指标，避免未使用的 metric 依赖在导入阶段直接报错。")
    add_paragraph(doc, "4. 额外安装 matplotlib、rouge_score、absl-py、sacrebleu，用于跑 sense_label 与 CoDWoE 自动评测。")

    add_paragraph(
        doc,
        "EACL 2024 仓库修改：",
    )
    add_paragraph(doc, "1. 将 DeepMistake、部分 WSI/LSCD 组件改为懒加载，避免未使用的可选依赖在导入阶段直接报错。")
    add_paragraph(doc, "2. 为 lemma.py 与 dataset.py 增加 .tsv 回退支持，以兼容当前官方数据。")
    add_paragraph(doc, "3. 修正 nordiachange_1.yaml 中的 grouping 名称，使其与当前上游数据中的 1929-1965 / 1970-2015 一致。")

    doc.add_heading("5. 未完全覆盖的部分与原因", level=1)
    add_paragraph(
        doc,
        "1. ACL 2023 已经补上了 CoDWoE English trial 这一层，但完整 WordNet / Oxford / 全量 CoDWoE 大规模评测尚未全部跑完。",
    )
    add_paragraph(
        doc,
        "2. EACL 2024 仓库原本准备先使用 testwug_en_111 做最小 benchmark，但该数据集下载地址在本机网络环境下无法访问 zenodo，因此最终切换到可通过 GitHub 获取的官方数据配置 nordiachange_1。",
    )
    add_paragraph(
        doc,
        "3. 当前机器以本地可运行、可提交、可解释为优先，因此在 benchmark 中使用了较小 checkpoint，而不是把全部官方模型与全部数据集都完整跑完。",
    )

    doc.add_heading("6. 最终结论", level=1)
    add_paragraph(
        doc,
        "如果按“老师要求给出 docx 说明文档，且要基本上复现论文”来判断，本次工作已经满足一个比较稳妥的提交标准：",
    )
    add_paragraph(doc, "1. 两个官方代码仓库都在本机真实跑通。")
    add_paragraph(doc, "2. ACL 2023 已经在真实语义变化数据上复现了核心方法链路，并新增了一层 CoDWoE English trial 金标 benchmark。")
    add_paragraph(doc, "3. EACL 2024 已经在官方 benchmark 框架 + 官方数据配置的真实子集上跑出了可记录结果。")
    add_paragraph(doc, "4. 所有输入、输出、脚本和结果文件都已经整理进 reproduction_pack 目录。")
    add_paragraph(
        doc,
        "因此，可以把本次结果表述为：完成了两篇论文的本地基本复现，其中 ACL 2023 同时覆盖了方法流程复现和一层带金标的定义生成评测，EACL 2024 侧重 benchmark 框架复现；尚未完全覆盖全部大规模实验。",
    )

    doc.add_heading("7. 关键文件位置", level=1)
    add_paragraph(doc, "总目录：/Users/sdf/Desktop/论文复现/reproduction_pack")
    add_paragraph(doc, "总索引：/Users/sdf/Desktop/论文复现/reproduction_pack/INDEX.md")
    add_paragraph(doc, "ACL 2023 目录：/Users/sdf/Desktop/论文复现/reproduction_pack/acl2023_interpretable_definition_generation")
    add_paragraph(doc, "EACL 2024 目录：/Users/sdf/Desktop/论文复现/reproduction_pack/eacl2024_semantic_change_benchmark")

    output = PACK / "论文复现说明.docx"
    root_output = ROOT / "论文复现说明.docx"
    doc.save(output)
    doc.save(root_output)
    print(output)
    print(root_output)


if __name__ == "__main__":
    main()
