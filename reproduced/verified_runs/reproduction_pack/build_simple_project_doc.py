#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path("/Users/sdf/Desktop/论文复现")
OUT = ROOT / "项目运行与文件说明.docx"


def set_font(run, east_asia: str = "Songti SC", size: float = 10.5, ascii_font: str = "Times New Roman") -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)


def add_paragraph(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    set_font(run)


def add_code(doc: Document, text: str) -> None:
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        run = p.add_run(line)
        set_font(run, east_asia="Courier New", size=9, ascii_font="Courier New")


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    set_font(run, size={1: 14, 2: 12, 3: 11}.get(level, 10.5))


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Pt(72)
    sec.right_margin = Pt(72)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("项目运行与文件说明")
    run.bold = True
    set_font(run, size=16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"目录：{ROOT}")
    set_font(run)

    add_paragraph(doc, "这份文档简单说明这个工作区里两个主要复现项目怎么运行，以及每个重要目录/文件大概是做什么的。")

    add_heading(doc, "1. 工作区整体结构", 1)
    top_items = [
        "LSCDBenchmark/：EACL 2024 语义变化 benchmark 框架，本地已经跑通过一个可验证子集。",
        "definition_modeling/：ACL 2023 定义生成项目，负责根据上下文生成词义定义。",
        "reproduction_pack/：把本地复现时用到的脚本、输入样例、输出结果整理在一起，最适合直接查看。",
        "papers/：其他论文和代码备份，属于资料区，不是这次最主要的两个复现入口。",
        "datasets/、acl2023_benchmarks/、acl2023_data/：数据文件或压缩包。",
        "runs/：部分运行中间结果和样例输出。",
        "codezips/、archives/：代码压缩包和归档文件。",
        ".conda/：本地 Conda 环境目录。",
    ]
    for item in top_items:
        add_paragraph(doc, item)

    add_heading(doc, "2. 代码怎么运行", 1)
    add_paragraph(doc, "当前最清晰的运行入口在 reproduction_pack/ 下面，已经整理成可直接执行的脚本。")

    add_heading(doc, "2.1 ACL 2023 定义生成项目", 2)
    add_paragraph(doc, "项目目录：/Users/sdf/Desktop/论文复现/definition_modeling")
    add_paragraph(doc, "配套说明：definition_modeling/README.md")
    add_paragraph(doc, "复现脚本目录：reproduction_pack/acl2023_interpretable_definition_generation/commands/")
    add_paragraph(doc, "建议运行方式：")
    add_code(
        doc,
        """cd /Users/sdf/Desktop/论文复现/reproduction_pack/acl2023_interpretable_definition_generation/commands
bash run_smoke_test.sh
bash run_semantic_change_pipeline.sh
bash run_codwoe_trial_benchmark.sh""",
    )
    add_paragraph(doc, "三个脚本的含义：")
    add_paragraph(doc, "run_smoke_test.sh：最小样例测试，先验证环境和模型能不能正常跑起来。")
    add_paragraph(doc, "run_semantic_change_pipeline.sh：真实语义变化样例主流程，会生成 definitions、embedding 和 sense labels。")
    add_paragraph(doc, "run_codwoe_trial_benchmark.sh：在 CoDWoE English trial 上做定义生成和简单自动评测。")

    add_heading(doc, "2.2 EACL 2024 语义变化 benchmark", 2)
    add_paragraph(doc, "项目目录：/Users/sdf/Desktop/论文复现/LSCDBenchmark")
    add_paragraph(doc, "配套说明：LSCDBenchmark/README.md")
    add_paragraph(doc, "复现脚本目录：reproduction_pack/eacl2024_semantic_change_benchmark/commands/")
    add_paragraph(doc, "建议运行方式：")
    add_code(
        doc,
        """cd /Users/sdf/Desktop/论文复现/reproduction_pack/eacl2024_semantic_change_benchmark/commands
bash run_verified_subset.sh""",
    )
    add_paragraph(doc, "这个脚本会调用 LSCDBenchmark 的主程序，在选定词集合上跑一遍已验证的 benchmark 配置。")

    add_heading(doc, "2.3 环境文件", 2)
    add_paragraph(doc, "definition_modeling/README.md：ACL 2023 项目的官方使用说明。")
    add_paragraph(doc, "LSCDBenchmark/README.md：EACL 2024 benchmark 项目的官方使用说明。")
    add_paragraph(doc, "LSCDBenchmark/requirements.txt：benchmark 项目依赖列表。")
    add_paragraph(doc, "papers/2025-semantic-change-discovery/environment.yml：另一个项目的 Conda 环境文件，可作参考，但不是这两个主项目的唯一入口。")

    add_heading(doc, "3. 每个重要文件/目录大概是什么意思", 1)
    detail_items = [
        "reproduction_pack/INDEX.md：整个复现包的总说明，先看这个最省时间。",
        "reproduction_pack/manifest.tsv：机器可读的复现摘要。",
        "reproduction_pack/build_docx_report.py：之前用于生成详细复现报告的脚本。",
        "reproduction_pack/acl2023_interpretable_definition_generation/README.md：ACL 2023 这部分复现内容的详细说明。",
        "reproduction_pack/acl2023_interpretable_definition_generation/artifacts/：ACL 2023 的输入、生成结果、评测结果。",
        "reproduction_pack/acl2023_interpretable_definition_generation/commands/prepare_semantic_change_subset.py：整理语义变化样例输入。",
        "reproduction_pack/acl2023_interpretable_definition_generation/commands/prepare_codwoe_trial.py：整理 CoDWoE trial 输入。",
        "reproduction_pack/acl2023_interpretable_definition_generation/commands/postprocess_generated_definitions.py：对生成定义做后处理，并抽取标签。",
        "reproduction_pack/eacl2024_semantic_change_benchmark/README.md：EACL 2024 这部分复现内容的详细说明。",
        "reproduction_pack/eacl2024_semantic_change_benchmark/artifacts/：EACL 2024 的配置、标签子集、预测结果和最终结果。",
        "reproduction_pack/eacl2024_semantic_change_benchmark/commands/run_verified_subset.sh：已验证跑通的 benchmark 命令。",
        "acl2023_benchmarks/：ACL 2023 用到的 benchmark 数据。",
        "definition_modeling/：原始定义生成代码仓库，核心模型代码在里面。",
        "LSCDBenchmark/：原始语义变化 benchmark 代码仓库，主入口通常是 main.py。",
    ]
    for item in detail_items:
        add_paragraph(doc, item)

    add_heading(doc, "4. 建议查看顺序", 1)
    for item in [
        "先看 reproduction_pack/INDEX.md，快速了解整个复现包。",
        "如果看 ACL 2023，就看 reproduction_pack/acl2023_interpretable_definition_generation/README.md，再看 commands/ 和 artifacts/。",
        "如果看 EACL 2024，就看 reproduction_pack/eacl2024_semantic_change_benchmark/README.md，再看 run_verified_subset.sh 和 artifacts/。",
        "如果需要回到原项目代码，再进入 definition_modeling/ 或 LSCDBenchmark/ 查看官方 README 和源码。",
    ]:
        add_paragraph(doc, item)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
